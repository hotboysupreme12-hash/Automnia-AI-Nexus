from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "dystopai_subscription_auth_launch_guide.docx"
EMAIL_DRAFT_PATH = OUT_DIR / "dystopai_subscription_auth_launch_guide_email_draft.txt"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths_inches)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = int(widths_inches[i] * 1440)
            cell.width = Inches(widths_inches[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_table_borders(table, color="D0D7DE", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(72, 82, 95)
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor(22, 27, 34)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.05

    small = styles.add_style("SmallMuted", 1)
    small.font.name = "Calibri"
    small.font.size = Pt(9)
    small.font.color.rgb = RGBColor(72, 82, 95)
    small.paragraph_format.space_after = Pt(4)


def add_header_footer(doc: Document) -> None:
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = "DystopAI Subscription, OAuth, and Authorization Launch Guide"
    paragraph.style = doc.styles["SmallMuted"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Prepared for DystopAI Core - June 4, 2026"
    paragraph.style = doc.styles["SmallMuted"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc: Document, text: str, style: str | None = None):
    return doc.add_paragraph(text, style=style)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbers(doc: Document, items: list[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(ind)
    for node in [start, num_fmt, lvl_text, lvl_jc, p_pr]:
        lvl.append(node)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.append(num_pr)
        p.add_run(item)


def add_callout(doc: Document, label: str, body: str, fill="F4F6F9", border="D0D7DE") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color=border)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], "E8EEF5")
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(11, 37, 69)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(text)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D8DEE4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    cell.paragraphs[0].text = ""
    lines = code.strip("\n").split("\n")
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.style = doc.styles["CodeBlock"]
        paragraph.add_run(line.rstrip())
    doc.add_paragraph()


def add_source_list(doc: Document) -> None:
    doc.add_heading("Source List", level=1)
    add_para(doc, "Official docs checked on June 4, 2026:", style="SmallMuted")
    sources = [
        "Stripe subscription webhooks: https://docs.stripe.com/billing/subscriptions/webhooks",
        "Stripe entitlements: https://docs.stripe.com/billing/entitlements",
        "Stripe customer portal: https://docs.stripe.com/billing/subscriptions/customer-portal",
        "Google OAuth for desktop apps: https://developers.google.com/identity/protocols/oauth2/native-app",
        "Clerk social connections: https://clerk.com/docs/guides/configure/auth-strategies/social-connections/overview",
        "Clerk production deployment: https://clerk.com/docs/guides/development/deployment/production",
        "Apple App Review Guidelines: https://developer.apple.com/app-store/review/guidelines/",
        "Sign in with Apple for web: https://developer.apple.com/help/account/capabilities/configure-sign-in-with-apple-for-the-web",
        "Google Play Payments policy: https://support.google.com/googleplay/android-developer/answer/9858738",
        "Paddle customer portal: https://developer.paddle.com/concepts/sell/customer-portal/",
        "Lemon Squeezy customer portal: https://docs.lemonsqueezy.com/help/online-store/customer-portal",
        "Lemon Squeezy merchant of record: https://docs.lemonsqueezy.com/help/payments/merchant-of-record",
    ]
    add_bullets(doc, sources)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    doc.add_paragraph("DystopAI Subscription, OAuth, and Authorization Launch Guide", style="Title")
    doc.add_paragraph(
        "A practical first implementation plan for selling the Electron/web app with Google login, Apple login, Stripe subscriptions, and server-side entitlement checks.",
        style="Subtitle",
    )
    add_callout(
        doc,
        "Bottom line",
        "For the first saleable version, ship direct-download desktop plus hosted web account pages. Use Clerk for Google/Apple/email login, Stripe Billing for subscriptions, and your own backend as the only authority for app access. The desktop app should only ask your backend whether the signed-in user is currently entitled.",
        fill="EAF3FF",
        border="B6D7F5",
    )

    doc.add_heading("1. What To Do First", level=1)
    add_numbers(
        doc,
        [
            "Pick the first sales channel: direct-download desktop plus web account portal. Avoid iOS App Store and Google Play until the account and entitlement system works.",
            "Buy or finalize the production domain, for example dystopai.com, and set DNS through Cloudflare or another DNS provider.",
            "Create production accounts for Clerk, Stripe, a Postgres database, and hosting. Add Apple Developer and Google Cloud once you configure Apple/Google login.",
            "Replace the current dev-token gate in this repo with real hosted account login and a server entitlement endpoint.",
            "Create Stripe products/prices and configure Checkout, Customer Portal, webhooks, and entitlements.",
            "Deploy the backend first, then wire the Electron app to open browser login, poll for completion, store the resulting app session securely, and check entitlements at launch.",
            "Run subscription-state tests before launch: active, trialing, past_due, canceled, unpaid, refunded, webhook retry, and offline grace.",
        ],
    )

    doc.add_heading("2. Recommended V1 Stack", level=1)
    add_table(
        doc,
        ["Layer", "Recommended choice", "Why"],
        [
            ["Identity", "Clerk", "Fast Google, Apple, email, passkeys, session management, production OAuth setup, and webhooks."],
            ["Billing", "Stripe Billing", "Best flexibility for subscriptions, Checkout, customer portal, webhooks, entitlements, coupons, invoices, and B2B expansion."],
            ["Database", "Postgres via Supabase or Neon", "Simple hosted relational store for users, Stripe customer IDs, subscription state, devices, and entitlement cache."],
            ["Backend", "Your existing Express/TypeScript server, deployed separately", "The current app already has Express. Keep the trusted subscription checks server-side."],
            ["Desktop auth", "External browser login plus device-style polling", "Works cleanly for Electron without embedding OAuth pages or putting provider secrets in the app."],
            ["Subscription management", "Stripe Customer Portal", "Users can update cards, change plans, download invoices, and cancel without you building billing UI."],
        ],
        [1.2, 1.6, 3.7],
    )

    add_callout(
        doc,
        "When to choose Paddle or Lemon Squeezy instead",
        "Choose Paddle or Lemon Squeezy if tax/VAT, merchant-of-record handling, and built-in customer billing management matter more than Stripe-level flexibility. Keep the same app architecture: auth provider plus billing webhooks plus your own entitlement API.",
        fill="FFF8E5",
        border="E5C76B",
    )

    doc.add_page_break()
    doc.add_heading("3. Accounts To Create", level=1)
    add_table(
        doc,
        ["Account", "Use it for", "First setup tasks"],
        [
            ["Domain/DNS", "Production URLs and OAuth redirect domains", "Own domain, add DNS records, create app.dystopai.com and api.dystopai.com."],
            ["Clerk", "User accounts and social login", "Create production instance, enable Google/Apple/email, set allowed redirect URLs, create webhook signing secret."],
            ["Google Cloud", "Google OAuth credentials", "Create OAuth consent screen, web OAuth client for Clerk, configure authorized domains, publish or verify if scopes require it."],
            ["Apple Developer", "Sign in with Apple and future macOS/iOS distribution", "Register App ID, Services ID, Sign in with Apple key, domain/return URLs."],
            ["Stripe", "Subscriptions and billing portal", "Activate account, create products/prices, customer portal, webhook endpoint, entitlements, test mode first."],
            ["Database", "Subscription truth cache", "Create Postgres project, run schema, keep service credentials only on backend."],
            ["Hosting", "Backend and web account portal", "Deploy server/API, set HTTPS, set environment variables, add logs and health check."],
            ["Transactional email", "Receipts beyond Stripe, support, security notices", "Use Resend/Postmark/SES later; Stripe handles payment receipts to start."],
            ["Code signing", "User trust for installers", "Windows Authenticode or Azure Trusted Signing; Apple Developer ID/notarization if shipping macOS."],
        ],
        [1.35, 1.75, 3.4],
    )

    doc.add_page_break()
    doc.add_heading("4. Architecture", level=1)
    add_para(doc, "The desktop app is not the authority. Your backend is. Stripe tells your backend what was paid. Clerk tells your backend who the user is. The app receives only a limited app session and an entitlement response.")
    add_code(
        doc,
        """
User -> Electron app
Electron app -> POST /api/desktop-auth/start
Backend -> returns browser login URL + one-time desktop session id
Electron app -> opens system browser
Browser -> Clerk login with Google / Apple / email
Web account page -> marks desktop session approved
Electron app -> polls /api/desktop-auth/poll/:id
Backend -> returns app session token
Electron app -> GET /api/me/entitlements
Backend -> reads DB/Stripe cache and returns allowed features
Electron app -> unlocks only features listed by backend
""",
    )

    doc.add_page_break()
    doc.add_heading("5. What Code You Need To Write", level=1)
    add_para(doc, "Minimum production code for V1:")
    add_bullets(
        doc,
        [
            "Auth/session backend routes for desktop login start, browser approval, polling, refresh, and logout.",
            "Stripe Checkout route to start a subscription.",
            "Stripe Customer Portal route for plan/payment management.",
            "Stripe webhook route that verifies webhook signatures and updates your DB.",
            "Entitlement route that returns the current plan/features to the app.",
            "Frontend auth state replacing the current local dev token gate.",
            "Electron secure token storage and app-launch entitlement check.",
            "Optional device registration if one subscription should only activate a limited number of machines.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("6. Database Schema", level=1)
    add_para(doc, "Use Postgres. Store provider IDs and Stripe IDs, but never store raw card data or OAuth provider secrets.")
    add_code(
        doc,
        """
create table app_users (
  id uuid primary key default gen_random_uuid(),
  clerk_user_id text unique not null,
  email text not null,
  stripe_customer_id text unique,
  created_at timestamptz not null default now()
);

create table subscriptions (
  id uuid primary key default gen_random_uuid(),
  user_id uuid not null references app_users(id) on delete cascade,
  stripe_subscription_id text unique not null,
  stripe_customer_id text not null,
  status text not null,
  price_id text,
  current_period_end timestamptz,
  cancel_at_period_end boolean not null default false,
  updated_at timestamptz not null default now()
);

create table entitlements (
  user_id uuid primary key references app_users(id) on delete cascade,
  plan text not null default 'free',
  features text[] not null default '{}',
  source text not null default 'stripe',
  expires_at timestamptz,
  updated_at timestamptz not null default now()
);

create table desktop_auth_sessions (
  id uuid primary key default gen_random_uuid(),
  code_hash text unique not null,
  clerk_user_id text,
  status text not null default 'pending',
  expires_at timestamptz not null,
  created_at timestamptz not null default now()
);

create table user_devices (
  id uuid primary key default gen_random_uuid(),
  user_id uuid not null references app_users(id) on delete cascade,
  device_fingerprint_hash text not null,
  device_name text,
  last_seen_at timestamptz not null default now(),
  revoked_at timestamptz,
  unique (user_id, device_fingerprint_hash)
);
""",
    )

    doc.add_heading("7. Backend Route Contracts", level=1)
    add_table(
        doc,
        ["Route", "Purpose", "Trust rule"],
        [
            ["POST /api/desktop-auth/start", "Create one-time desktop login session and return browser login URL.", "No user trust yet. Generate random code and expiry."],
            ["POST /api/desktop-auth/approve", "Called by signed-in web page after Clerk auth.", "Requires valid Clerk session on web."],
            ["GET /api/desktop-auth/poll/:id", "Desktop polls until approved.", "Return app token only once, then expire session."],
            ["POST /api/billing/checkout", "Create Stripe Checkout subscription session.", "Requires authenticated user."],
            ["POST /api/billing/portal", "Create Stripe Customer Portal session.", "Requires authenticated user with Stripe customer ID."],
            ["POST /api/stripe/webhook", "Receive Stripe subscription and entitlement changes.", "Must use raw body and verify Stripe signature."],
            ["GET /api/me/entitlements", "Return current feature access for the app.", "Requires app session. Compute from DB cache, reconcile if stale."],
        ],
        [2.05, 2.65, 1.8],
    )

    doc.add_heading("8. Starter Backend Skeleton", level=1)
    add_para(doc, "This is a skeleton, not a paste-and-launch production server. It shows the shape of the code you add to your existing Express backend.")
    add_code(
        doc,
        """
import express from 'express'
import Stripe from 'stripe'
import { ClerkExpressRequireAuth } from '@clerk/clerk-sdk-node'

const app = express()
const stripe = new Stripe(process.env.STRIPE_SECRET_KEY!)

app.post(
  '/api/stripe/webhook',
  express.raw({ type: 'application/json' }),
  async (req, res) => {
    const sig = req.header('stripe-signature')
    let event: Stripe.Event

    try {
      event = stripe.webhooks.constructEvent(
        req.body,
        sig!,
        process.env.STRIPE_WEBHOOK_SECRET!,
      )
    } catch {
      return res.status(400).send('Invalid webhook signature')
    }

    if (event.type === 'checkout.session.completed') {
      // Find user from session metadata.clerkUserId.
      // Store stripe_customer_id and subscription id.
    }

    if (event.type === 'customer.subscription.updated' ||
        event.type === 'customer.subscription.deleted') {
      // Upsert subscription status and current_period_end.
      // Recompute entitlements for that user.
    }

    if (event.type === 'entitlements.active_entitlement_summary.updated') {
      // Map Stripe entitlement lookup_keys into your local feature list.
    }

    return res.json({ received: true })
  },
)

app.use(express.json())

app.post('/api/billing/checkout', ClerkExpressRequireAuth(), async (req, res) => {
  const clerkUserId = req.auth.userId
  const user = await findOrCreateUser(clerkUserId)

  const session = await stripe.checkout.sessions.create({
    mode: 'subscription',
    customer: user.stripe_customer_id ?? undefined,
    customer_email: user.stripe_customer_id ? undefined : user.email,
    line_items: [{ price: process.env.STRIPE_PRICE_PRO_MONTHLY!, quantity: 1 }],
    success_url: `${process.env.APP_URL}/billing/success`,
    cancel_url: `${process.env.APP_URL}/billing/cancel`,
    metadata: { clerkUserId },
  })

  res.json({ url: session.url })
})

app.post('/api/billing/portal', ClerkExpressRequireAuth(), async (req, res) => {
  const user = await getUserByClerkId(req.auth.userId)
  const portal = await stripe.billingPortal.sessions.create({
    customer: user.stripe_customer_id,
    return_url: `${process.env.APP_URL}/account`,
  })
  res.json({ url: portal.url })
})
""",
    )

    doc.add_heading("9. Frontend Changes In This Repo", level=1)
    add_para(doc, "The current app has a local token gate in src/context/AuthContext.tsx and src/components/auth/LoginModal.tsx. That is fine for development, but it cannot protect a paid app.")
    add_bullets(
        doc,
        [
            "Replace the password/token modal with buttons: Continue with Google, Continue with Apple, Continue with Email.",
            "For Electron, button click calls /api/desktop-auth/start and opens the returned URL in the system browser.",
            "Show a waiting state while the app polls /api/desktop-auth/poll/:id.",
            "After success, store the app session in OS secure storage. Do not use plain localStorage for production desktop sessions.",
            "At app launch, call /api/me/entitlements. If plan is free or inactive, keep the app usable but lock premium actions.",
            "Add Manage subscription button that calls /api/billing/portal and opens the Stripe portal URL.",
        ],
    )
    add_code(
        doc,
        """
type EntitlementsResponse = {
  active: boolean
  plan: 'free' | 'pro' | 'team'
  features: string[]
  expiresAt?: string
  graceUntil?: string
}

async function requireEntitlement(feature: string) {
  const res = await fetch('/api/me/entitlements', {
    headers: { Authorization: `Bearer ${sessionToken}` },
  })
  const data = await res.json() as EntitlementsResponse
  return data.active && data.features.includes(feature)
}
""",
    )

    doc.add_heading("10. Entitlement Rules", level=1)
    add_table(
        doc,
        ["Plan", "Features", "Server behavior"],
        [
            ["Free", "Basic local UI, account page, docs, purchase screen", "No paid model/runtime features. Good for onboarding."],
            ["Pro", "Premium agents, model providers, updates, limited devices", "Return active true while subscription is active/trialing or inside grace window."],
            ["Team", "Multiple users, shared billing, device seats, admin portal", "Add organization-level entitlements later after individual Pro works."],
        ],
        [1.2, 2.4, 2.9],
    )
    add_callout(
        doc,
        "Offline grace",
        "Cache a signed entitlement token for 24 to 72 hours so a paying user can open the desktop app during a temporary API outage. Do not cache forever. Re-check as soon as the backend is reachable.",
        fill="F4F6F9",
        border="D0D7DE",
    )

    doc.add_heading("11. Stripe Setup Checklist", level=1)
    add_numbers(
        doc,
        [
            "Create products: DystopAI Pro Monthly, DystopAI Pro Annual, optionally Team.",
            "Create recurring prices and record the price IDs in backend environment variables.",
            "Configure Customer Portal to allow payment method updates, invoice downloads, cancellation, and plan changes.",
            "Create entitlement features in Stripe, such as premium_agents, cloud_sync, priority_runtime, team_seats.",
            "Attach features to Stripe products.",
            "Create webhook endpoint: https://api.dystopai.com/api/stripe/webhook.",
            "Listen for checkout.session.completed, invoice.paid, invoice.payment_failed, customer.subscription.updated, customer.subscription.deleted, and entitlements.active_entitlement_summary.updated.",
            "Use Stripe CLI in test mode to simulate renewal, payment failure, cancellation, and entitlement update events.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("12. OAuth Setup Checklist", level=1)
    add_para(doc, "With Clerk, Google and Apple OAuth happen on the hosted web auth surface. Your desktop app should open a real external browser instead of embedding the provider login screen.")
    add_table(
        doc,
        ["Provider", "Account setup", "Important detail"],
        [
            ["Google", "Google Cloud OAuth consent screen plus web client credentials in Clerk.", "Desktop apps cannot keep client secrets. If you directly use Google OAuth, use Authorization Code with PKCE."],
            ["Apple", "Apple Developer App ID, Services ID, Sign in with Apple private key, domains, return URLs.", "If an Apple-platform app uses social login for primary account auth, Apple requires an equivalent privacy-focused login option."],
            ["Email", "Enable Clerk email/password, magic link, or passkeys.", "A first-party email login is useful fallback when users do not want Google or Apple."],
        ],
        [1.15, 2.65, 2.7],
    )

    doc.add_heading("13. App Store And Play Store Rules", level=1)
    add_callout(
        doc,
        "Direct desktop/web is simpler",
        "If you sell the Windows/macOS/Linux desktop installer from your own website, Stripe/Paddle/Lemon Squeezy are normal choices. If you sell digital features inside iOS App Store or Google Play apps, you must handle each store's in-app purchase policy before launch.",
        fill="FFF8E5",
        border="E5C76B",
    )
    add_bullets(
        doc,
        [
            "Apple App Store: subscriptions must deliver ongoing value and work across the user's devices where the app is available. In-app digital purchases generally trigger App Store purchase rules.",
            "Apple login: Apple guideline 4.8 requires apps using third-party/social login for the user's primary account to offer an equivalent privacy-focused login service, with listed exceptions.",
            "Google Play: Play-distributed apps requiring payment for in-app digital features, app functionality, subscriptions, or cloud software generally must use Google Play Billing unless a policy exception applies.",
            "Practical launch path: do direct desktop/web first. Add mobile store billing later only if you build real mobile apps.",
        ],
    )

    doc.add_heading("14. Security Rules", level=1)
    add_bullets(
        doc,
        [
            "Never place Stripe secret keys, Clerk secret keys, Apple private keys, or database credentials in the Electron client.",
            "Never trust subscription state sent by the app. The backend computes it from webhooks and reconciliation.",
            "Verify every Stripe webhook signature against the raw request body.",
            "Use short-lived app access tokens plus refresh tokens. Revoke refresh tokens on logout, refund fraud, or device revocation.",
            "Hash device fingerprints. Do not store unnecessary hardware identifiers in plaintext.",
            "Use HTTPS only in production. Add CSP on the web auth/account portal.",
            "Add audit logs for subscription changes, device activation, account deletion, and admin actions.",
            "Treat Electron client checks as user experience gates, not strong anti-piracy. Keep valuable server-backed services behind backend authorization.",
        ],
    )

    doc.add_heading("15. Launch Test Plan", level=1)
    add_table(
        doc,
        ["Test", "Expected result"],
        [
            ["New user signs in with Google", "Clerk user exists, local app gets app session, user has free entitlement."],
            ["New user signs in with Apple", "Private relay email works, account is linked, user can continue."],
            ["User buys Pro monthly", "Stripe Checkout completes, webhook stores customer/subscription, app unlocks Pro."],
            ["Payment fails", "Subscription becomes past_due; app shows warning and grace state."],
            ["Subscription cancels", "At period end, entitlement is revoked and paid features lock."],
            ["Webhook delivery delayed", "Nightly reconciliation detects mismatch and corrects DB."],
            ["Backend offline", "Cached signed entitlement works only inside grace window."],
            ["Device limit exceeded", "App shows device management path instead of silently failing."],
        ],
        [2.2, 4.3],
    )

    doc.add_page_break()
    doc.add_heading("16. First Seven Days Plan", level=1)
    add_table(
        doc,
        ["Day", "Work"],
        [
            ["1", "Choose final stack, create Clerk/Stripe/database/hosting accounts, set production domain plan."],
            ["2", "Create DB schema, build user sync from Clerk, deploy backend health endpoint."],
            ["3", "Build Stripe Checkout, Customer Portal, webhook verification, and local subscription cache."],
            ["4", "Replace dev token UI with desktop browser login and polling flow."],
            ["5", "Add /api/me/entitlements, feature gates, offline signed entitlement cache."],
            ["6", "Test all Stripe states with test cards and Stripe CLI. Add logging and error screens."],
            ["7", "Prepare launch pages: pricing, terms, privacy, support, refund/cancel policy, installer download."],
        ],
        [0.8, 5.7],
    )

    doc.add_page_break()
    doc.add_heading("17. Environment Variables", level=1)
    add_code(
        doc,
        """
APP_URL=https://app.dystopai.com
API_URL=https://api.dystopai.com
DATABASE_URL=postgres://...

CLERK_SECRET_KEY=sk_live_...
CLERK_WEBHOOK_SECRET=whsec_...

STRIPE_SECRET_KEY=sk_live_...
STRIPE_WEBHOOK_SECRET=whsec_...
STRIPE_PRICE_PRO_MONTHLY=price_...
STRIPE_PRICE_PRO_ANNUAL=price_...

APP_SESSION_JWT_SECRET=generate-a-long-random-secret
ENTITLEMENT_SIGNING_SECRET=generate-a-different-long-random-secret
""",
    )

    doc.add_heading("18. Definition Of Done Before Selling", level=1)
    add_bullets(
        doc,
        [
            "A user can sign up with Google, Apple, or email.",
            "A user can subscribe from a hosted pricing/account page.",
            "A user can open the desktop app and see paid features unlocked only after active entitlement.",
            "A user can manage/cancel their subscription through a hosted portal.",
            "Your backend survives webhook retries and can reconcile with Stripe.",
            "Your privacy policy, terms, support contact, and account deletion path exist.",
            "Your production keys are not in the repo, app bundle, public logs, or installer.",
            "The direct-download installer is code signed or at least clearly identified while you complete code signing.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("19. What I Would Build First In This Repo", level=1)
    add_numbers(
        doc,
        [
            "Add a separate production backend deployment instead of relying only on the local Electron-bundled server.",
            "Create new server modules: authDesktop.ts, billingStripe.ts, entitlementService.ts, deviceService.ts.",
            "Replace src/context/AuthContext.tsx with session/entitlement state instead of control-center-token.",
            "Replace src/components/auth/LoginModal.tsx with browser login actions and a waiting state.",
            "Add a Settings/Billing surface with current plan, renewal date, device list, and Manage subscription button.",
            "Keep existing provider API-key auth separate from app-account auth. Provider keys let the runtime call models; subscription auth controls whether the customer can use premium app features.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("20. Final Notes", level=1)
    add_para(doc, "This guide is technical implementation guidance, not legal, tax, or app-review legal advice. Before public launch, confirm tax, privacy, refund, and platform policy obligations for the countries and stores where you sell.")
    add_para(doc, "The strongest commercial control is not a hidden local license check. It is a server-backed account, webhook-updated entitlements, secure session handling, and premium services that require backend authorization.")

    add_source_list(doc)

    doc.save(DOCX_PATH)

    EMAIL_DRAFT_PATH.write_text(
        """To: hotboysupreme12@gmail.com
Subject: DystopAI Subscription/Auth Launch Guide

Attached is the DystopAI Subscription, OAuth, and Authorization Launch Guide PDF.

It covers the first steps, accounts to create, recommended stack, Stripe/Clerk setup, OAuth rules, backend route contracts, database schema, starter code, security rules, and a seven-day implementation plan.

Attachment:
dystopai_subscription_auth_launch_guide.pdf
""",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
    print(EMAIL_DRAFT_PATH)
